"""
Extractor module: Extract text from PDFs (with 2-column layout detection),
DOCX, TXT files, with OCR fallback for scanned PDFs.
"""
# pyright: reportMissingImports=false, reportMissingModuleSource=false

import io
import logging
from typing import Tuple, Optional
import pdfplumber
from pdf2image import convert_from_bytes
import pytesseract
from docx import Document
import chardet
from langdetect import detect, LangDetectException

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract and clean text from various file formats"""
    
    @staticmethod
    def extract_pdf(content: bytes) -> Tuple[str, str, bool]:
        """
        Extract text from PDF with 2-column layout detection and OCR fallback.
        Returns: (text, quality, was_ocr_used)
        """
        try:
            try:
                pdf = pdfplumber.open(io.BytesIO(content))
            except Exception as e:
                logger.warning(f"Failed to open PDF with pdfplumber: {e}")
                return TextExtractor._ocr_pdf_with_quality(content)
            
            all_text = []
            was_ocr = False
            
            for page_num, page in enumerate(pdf.pages):
                page_text = TextExtractor._extract_page_with_layout(page)
                if page_text:
                    all_text.append(page_text)
                else:
                    logger.debug(f"No text extracted from page {page_num + 1}, attempting OCR")
            
            combined_text = "\n".join(all_text)
            pdf.close()
            
            # Check if we got meaningful text
            if len(combined_text.strip()) < 100:
                # Fallback to OCR
                logger.info("Extracted text is too short, falling back to OCR")
                return TextExtractor._ocr_pdf_with_quality(content)
            
            return combined_text, "good", was_ocr
        except Exception as e:
            logger.error(f"Error extracting PDF: {e}")
            # Fallback to OCR on error
            return TextExtractor._ocr_pdf_with_quality(content)
    
    @staticmethod
    def _extract_page_with_layout(page) -> str:
        """
        Extract text from a PDF page, detecting 2-column layout.
        If 2-column: extract left column, then right column.
        Otherwise: use default layout extraction.
        """
        try:
            words = page.extract_words(use_text_flow=False)
            if not words:
                # Try with text_flow enabled as fallback
                try:
                    return page.extract_text(layout=True) or ""
                except:
                    return ""
            
            # Get x0 coordinates to detect column layout
            x_coords = [w['x0'] for w in words]
            x_coords_sorted = sorted(set(x_coords))  # Remove duplicates for better gap detection
            
            # Simple 2-column detection: look for gap in x distribution
            if len(x_coords_sorted) > 10:
                gap = TextExtractor._find_column_gap(x_coords_sorted)
                if gap is not None:
                    # Split into left and right by x coordinate
                    left_words = [w for w in words if w['x0'] < gap]
                    right_words = [w for w in words if w['x0'] >= gap]
                    
                    # Validate split
                    if left_words and right_words:
                        # Sort each column by top position, then left position
                        left_text = " ".join([w['text'] for w in sorted(left_words, key=lambda x: (x['top'], x['x0']))])
                        right_text = " ".join([w['text'] for w in sorted(right_words, key=lambda x: (x['top'], x['x0']))])
                        
                        return f"{left_text}\n{right_text}"
            
            # Default: use layout extraction
            return page.extract_text(layout=True) or ""
        except Exception as e:
            logger.debug(f"Error in page layout extraction: {e}")
            try:
                return page.extract_text() or ""
            except:
                return ""
    
    @staticmethod
    def _find_column_gap(sorted_x_coords: list) -> Optional[float]:
        """Find natural gap in x coordinates indicating column separation"""
        if len(sorted_x_coords) < 2:
            return None
            
        # Check for gap in middle ranges
        gaps = []
        for i in range(len(sorted_x_coords) - 1):
            gap_size = sorted_x_coords[i + 1] - sorted_x_coords[i]
            gaps.append((gap_size, sorted_x_coords[i]))
        
        # Find largest gap in reasonable range (100-600px for flexibility)
        significant_gaps = [(g, x) for g, x in gaps if 100 < g < 600]
        if significant_gaps:
            gap_size, gap_x = max(significant_gaps, key=lambda x: x[0])
            return gap_x + gap_size / 2
        
        # Fallback: use median x-coordinate if no clear gap
        if sorted_x_coords:
            return sorted_x_coords[len(sorted_x_coords) // 2]
        
        return None
    
    @staticmethod
    def _ocr_pdf_with_quality(content: bytes) -> Tuple[str, str, bool]:
        """OCR PDF and return with quality marker"""
        text, quality = TextExtractor._ocr_pdf(content)
        return text, quality, True
    
    @staticmethod
    def _ocr_pdf(content: bytes) -> Tuple[str, str]:
        """Fallback: OCR scanned PDF using Tesseract"""
        try:
            images = convert_from_bytes(content, dpi=300)
            all_text = []
            for page_num, img in enumerate(images):
                try:
                    text = pytesseract.image_to_string(img)
                    if text.strip():
                        all_text.append(text)
                except Exception as e:
                    logger.warning(f"OCR failed for page {page_num + 1}: {e}")
            
            combined = "\n".join(all_text)
            quality = "ocr_used" if combined.strip() else "poor"
            return combined, quality
        except Exception as e:
            logger.error(f"OCR PDF failed completely: {e}")
            return "", "poor"
    
    @staticmethod
    def extract_docx(content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(io.BytesIO(content))
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX: {e}")
            return ""
    
    @staticmethod
    def extract_txt(content: bytes) -> str:
        """Extract text from TXT file with encoding detection"""
        try:
            # Try UTF-8 first
            return content.decode('utf-8')
        except UnicodeDecodeError:
            # Detect encoding
            try:
                detected = chardet.detect(content)
                encoding = detected.get('encoding', 'latin-1')
                if not encoding:
                    encoding = 'latin-1'
                return content.decode(encoding, errors='replace')
            except Exception as e:
                logger.warning(f"Encoding detection failed: {e}")
                return content.decode('latin-1', errors='replace')
    
    @staticmethod
    def extract_text(filename: str, content: bytes) -> Tuple[str, str, Optional[str]]:
        """
        Main method: extract text from any supported format.
        Returns: (text, quality, language)
        """
        if not filename or not content:
            logger.warning("Empty filename or content provided")
            return "", "poor", None
            
        ext = filename.lower().split('.')[-1]
        
        if ext == 'pdf':
            text, quality, _ = TextExtractor.extract_pdf(content)
        elif ext in ['docx', 'doc']:
            text = TextExtractor.extract_docx(content)
            quality = "good" if len(text.strip()) > 100 else "poor"
        elif ext == 'txt':
            text = TextExtractor.extract_txt(content)
            quality = "good" if len(text.strip()) > 100 else "poor"
        else:
            logger.warning(f"Unsupported file format: {ext}")
            return "", "poor", None
        
        # Detect language
        language = TextExtractor._detect_language(text)
        
        return text, quality, language
    
    @staticmethod
    def _detect_language(text: str) -> Optional[str]:
        """Detect language using langdetect"""
        try:
            if len(text.strip()) < 50:
                return None
            lang = detect(text[:500])
            return lang
        except LangDetectException:
            logger.debug("Language detection failed")
            return None
        except Exception as e:
            logger.error(f"Unexpected error in language detection: {e}")
            return None
